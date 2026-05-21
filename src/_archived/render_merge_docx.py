"""
render_merge_docx.py — 把已填好内容的合并 JSON 渲染成 Word (.docx) 报告。

颜色继承：
  - merged JSON 中 aligned.groups 各品种若含 prev_annotation 字段，
    则「备注」列显示继承的批注文本，并按 color_palette 对整行着底色。
  - color_palette 由 Claude Code 在合并分析阶段填写：
    { "颜色名": "RRGGBB", ... }，如 {"夕阳红": "C04020", "鹅黄": "FFEC8B"}
  - 若 color_name 在 palette 里找不到，跳过着色（不报错）。
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── 视觉规范（固化）──────────────────────────────────────────────────────────
FONT_BODY = "Microsoft YaHei"
FONT_MONO = "Courier New"

COLOR_UP      = RGBColor(0xFF, 0x00, 0x00)
COLOR_DOWN    = RGBColor(0x00, 0x00, 0x8B)
COLOR_ANOMALY = RGBColor(0xFF, 0x00, 0x00)

HEX_HEADER   = "D9D9D9"
HEX_HIST_ROW = "F5F5F5"   # 历史时段行底色（浅灰）

# 变化评级文字颜色
RATING_COLORS = {
    "大加强": COLOR_UP,
    "加强":   RGBColor(0xCC, 0x00, 0x00),
    "减弱":   RGBColor(0x00, 0x33, 0x88),
    "大减弱": COLOR_DOWN,
}

# 列定义（合并报告：「日期」→「时段」；末列「备注」）
MERGE_COLS: list[tuple[str, float]] = [
    ("名称",       2.5),
    ("时段",       2.0),
    ("今涨幅",     1.5),
    ("昨涨幅",     1.5),
    ("涨跌幅差值", 1.8),
    ("成交额环比", 1.8),
    ("归类特征",   1.5),
    ("归类",       1.8),
    ("符合情况",   1.8),
    ("原因分析",   5.5),
    ("备注",       3.0),
]

CAT_ORDER = ["持续强化", "强反转", "反包修复", "连续杀跌"]
CAT_ROMAN = {"持续强化": "一", "强反转": "二", "反包修复": "三", "连续杀跌": "四"}


# ── 底层工具（与 render_docx.py 保持一致）────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_row_bg(row, hex_color: str) -> None:
    for cell in row.cells:
        _set_cell_bg(cell, hex_color)


def _cell_write(cell, text: str, *,
                bold: bool = False,
                color: RGBColor | None = None,
                size: int = 9,
                align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    para = cell.paragraphs[0]
    para.alignment = align
    r = para.add_run(str(text))
    r.font.name = FONT_BODY
    r.font.size = Pt(size)
    if bold: r.bold = True
    if color: r.font.color.rgb = color


def _cell_write_analysis(cell, text: str, size: int = 8) -> None:
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if not text:
        r = para.add_run("")
        r.font.name = FONT_BODY
        r.font.size = Pt(size)
        return
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        m = re.match(r"\*\*([^*]+)\*\*", part)
        if m:
            inner = m.group(1)
            r = para.add_run(inner)
            r.font.name = FONT_BODY
            r.font.size = Pt(size)
            r.bold = True
            if inner == "异常":
                r.font.color.rgb = COLOR_ANOMALY
        elif part:
            r = para.add_run(part)
            r.font.name = FONT_BODY
            r.font.size = Pt(size)


def _pct_str(v: float) -> str:
    sign = "+" if v > 0 else ""
    return f"{sign}{v * 100:.2f}%"


def _pct_color(v: float) -> RGBColor | None:
    if v > 0: return COLOR_UP
    if v < 0: return COLOR_DOWN
    return None


def _para_text(doc: Document, text: str, *,
               bold: bool = False, size: int = 10,
               color: RGBColor | None = None,
               align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.name = FONT_BODY
    r.font.size = Pt(size)
    r.bold = bold
    if color: r.font.color.rgb = color


def _setup_landscape(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Cm(0.8))


def _make_table(doc: Document, col_defs: list) -> any:
    table = doc.add_table(rows=1, cols=len(col_defs))
    table.style = "Table Grid"
    for i, (_, w) in enumerate(col_defs):
        for cell in table.column_cells(i):
            cell.width = Cm(w)
    hrow = table.rows[0]
    _set_row_bg(hrow, HEX_HEADER)
    for i, (hdr, _) in enumerate(col_defs):
        _cell_write(hrow.cells[i], hdr, bold=True, size=9)
    return table


def _fill_data_row(cells, it: dict, label: str, *,
                   is_hist: bool = False,
                   note_text: str = "") -> None:
    """填写一行数据行（单时段 + 合并均用此函数）。"""
    size = 8 if is_hist else 9
    _cell_write(cells[0], it["name"],  bold=not is_hist, size=size)
    _cell_write(cells[1], label, size=size,
                color=RGBColor(0x66, 0x66, 0x66) if is_hist else None)
    _cell_write(cells[2], _pct_str(it["today_pct"]),
                color=_pct_color(it["today_pct"]), size=size)
    _cell_write(cells[3], _pct_str(it["yest_pct"]),
                color=_pct_color(it["yest_pct"]), size=size)
    _cell_write(cells[4], _pct_str(it["pct_diff"]),
                color=_pct_color(it["pct_diff"]), size=size)
    _cell_write(cells[5], _pct_str(it["volume_ratio"]),
                color=_pct_color(it["volume_ratio"]), size=size)
    _cell_write(cells[6], it.get("feature",    ""), size=size)
    _cell_write(cells[7], it.get("category",   ""), size=size)
    _cell_write(cells[8], it.get("compliance", ""), size=size)
    _cell_write_analysis(cells[9], it.get("analysis", ""), size=7 if is_hist else 8)
    _cell_write(cells[10], note_text, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)


# ── 主渲染函数 ────────────────────────────────────────────────────────────────

def render_merge(merged_json_path: str, output_path: str) -> Path:
    data = json.loads(Path(merged_json_path).read_text(encoding="utf-8"))

    aligned         = data["aligned"]
    snapshots       = data["snapshots_summary"]
    rating_cands    = data["rating_candidates"]
    cat_summaries   = data["category_summaries"]
    prev_annotations = data.get("prev_annotations", {})
    color_palette   = data.get("color_palette", {})
    macro           = data["macro"]
    resonance       = aligned.get("resonance")

    doc = Document()
    _setup_landscape(doc)

    # ── 标题 ──────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"{aligned['today_date']} 多时段合并研判报告")
    r.font.name = FONT_BODY
    r.font.size = Pt(18)
    r.bold = True

    # 时段列表
    labels_str = " ← ".join(s["session_label"] for s in snapshots)
    _para_text(doc, f"本次合并时段：{labels_str}", size=9,
               color=RGBColor(0x66, 0x66, 0x66), align=WD_ALIGN_PARAGRAPH.CENTER)

    # 共振预警
    if resonance == "up":
        _para_text(doc, "🚨 共振起爆：超70%品种上涨 🚨",
                   bold=True, size=13, color=COLOR_UP, align=WD_ALIGN_PARAGRAPH.CENTER)
    elif resonance == "down":
        _para_text(doc, "⚠️ 共振杀跌：超70%品种下跌 ⚠️",
                   bold=True, size=13, color=COLOR_DOWN, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── 第一部分：四象限剖析 ──────────────────────────────────────────────────
    h = doc.add_heading("第一部分 · 数据驱动的四象限剖析（含跨时段堆叠）", level=1)
    for run in h.runs: run.font.name = FONT_BODY

    stats = aligned["stats"]

    for cat in CAT_ORDER:
        group = aligned["groups"][cat]
        st    = stats[cat]

        h2 = doc.add_heading(
            f"{CAT_ROMAN[cat]}、{cat}类（T0 共 {st['count']} 个，占 {st['pct']}%）",
            level=2,
        )
        for run in h2.runs: run.font.name = FONT_BODY

        table = _make_table(doc, MERGE_COLS)

        for it in group:
            # ① 历史时段行（逆序：最旧 → 较新）
            for h_snap in reversed(it.get("timeline", [])):
                row   = table.add_row()
                cells = row.cells
                _set_row_bg(row, HEX_HIST_ROW)
                _fill_data_row(cells, h_snap, h_snap.get("label", ""),
                               is_hist=True, note_text="")

            # ② T0 行（带批注颜色继承）
            ann        = it.get("prev_annotation", {})
            note_text  = ann.get("note_text", "") if ann else ""
            color_name = ann.get("color_name", "") if ann else ""
            row_hex    = color_palette.get(color_name, "") if color_name else ""

            row   = table.add_row()
            cells = row.cells
            if row_hex:
                _set_row_bg(row, row_hex)
            else:
                pass  # 正常白色背景

            _fill_data_row(cells, it, aligned["session_label"],
                           is_hist=False, note_text=note_text)

        # 分类小结
        summary = cat_summaries.get(cat, "")
        if summary:
            p = doc.add_paragraph()
            r = p.add_run("【本分类分析小结】\n")
            r.font.name = FONT_BODY
            r.font.size = Pt(9)
            r.bold = True
            p2 = doc.add_paragraph(style="Normal")
            for part in re.split(r"(\*\*[^*]+\*\*)", summary):
                m = re.match(r"\*\*([^*]+)\*\*", part)
                if m:
                    inner = m.group(1)
                    run = p2.add_run(inner)
                    run.font.name = FONT_BODY
                    run.font.size = Pt(9)
                    run.bold = True
                    if inner == "异常":
                        run.font.color.rgb = COLOR_ANOMALY
                elif part:
                    run = p2.add_run(part)
                    run.font.name = FONT_BODY
                    run.font.size = Pt(9)

    # 独特异象
    if data.get("unique_anomaly_analysis"):
        h2 = doc.add_heading("独特异象追踪分析", level=2)
        for run in h2.runs: run.font.name = FONT_BODY
        p = doc.add_paragraph(style="Normal")
        for part in re.split(r"(\*\*[^*]+\*\*)", data["unique_anomaly_analysis"]):
            m = re.match(r"\*\*([^*]+)\*\*", part)
            if m:
                inner = m.group(1)
                r = p.add_run(inner)
                r.font.name = FONT_BODY
                r.font.size = Pt(9)
                r.bold = True
                if inner == "异常": r.font.color.rgb = COLOR_ANOMALY
            elif part:
                r = p.add_run(part)
                r.font.name = FONT_BODY
                r.font.size = Pt(9)

    # ── 跨日追踪表 ────────────────────────────────────────────────────────────
    h2 = doc.add_heading("跨日追踪表（昨日收盘 → T0 当前）", level=2)
    for run in h2.runs: run.font.name = FONT_BODY

    track_cols = [
        ("品种", 2.5), ("昨日特征", 2.0), ("昨日归类", 2.0),
        ("当前特征", 2.0), ("当前归类", 2.0), ("变化评级", 1.8), ("理由", 5.0),
    ]
    tt = _make_table(doc, track_cols)
    for c in rating_cands:
        row   = tt.add_row()
        cells = row.cells
        rating = c.get("rating", "")
        _cell_write(cells[0], c.get("name", ""), size=9)
        _cell_write(cells[1], c.get("yesterday_feature", ""), size=9)
        _cell_write(cells[2], c.get("yesterday_category", ""), size=9)
        _cell_write(cells[3], c.get("current_feature", ""),   size=9)
        _cell_write(cells[4], c.get("current_category", ""),  size=9)
        _cell_write(cells[5], rating, size=9,
                    color=RATING_COLORS.get(rating),
                    bold=rating in ("大加强", "大减弱"))
        _cell_write(cells[6], c.get("reason", ""), size=9,
                    align=WD_ALIGN_PARAGRAPH.LEFT)

    # ── 第二部分：宏观研判 ────────────────────────────────────────────────────
    h = doc.add_heading("第二部分 · 宏观研判", level=1)
    for run in h.runs: run.font.name = FONT_BODY

    def _section(doc, title, content_fn):
        h3 = doc.add_heading(title, level=3)
        for run in h3.runs: run.font.name = FONT_BODY
        content_fn()

    def _add_text_block(text: str):
        p = doc.add_paragraph(style="Normal")
        for part in re.split(r"(\*\*[^*]+\*\*)", text or ""):
            m = re.match(r"\*\*([^*]+)\*\*", part)
            if m:
                inner = m.group(1)
                r = p.add_run(inner)
                r.font.name = FONT_BODY
                r.font.size = Pt(9)
                r.bold = True
                if inner == "异常": r.font.color.rgb = COLOR_ANOMALY
            elif part:
                r = p.add_run(part)
                r.font.name = FONT_BODY
                r.font.size = Pt(9)

    # 1. 全景图
    def _panorama():
        p = doc.add_paragraph(style="Normal")
        r = p.add_run(macro.get("panorama_headline", ""))
        r.font.name = FONT_BODY
        r.font.size = Pt(11)
        r.bold = True
        for para_text in macro.get("panorama_paragraphs", []):
            if para_text:
                _add_text_block(para_text)
    _section(doc, "1. 市场资金全景图", _panorama)

    # 2. 关键异动板块
    def _key_movers():
        for mv in macro.get("key_movers", []):
            if not mv.get("sector"): continue
            h4 = doc.add_heading(f"▸ {mv['sector']}", level=4)
            for run in h4.runs: run.font.name = FONT_BODY
            for label, key in [("现象", "phenomenon"), ("动因", "driver"), ("推演", "scenario")]:
                p = doc.add_paragraph(style="Normal")
                r = p.add_run(f"{label}：")
                r.font.name = FONT_BODY
                r.font.size = Pt(9)
                r.bold = True
                r2 = p.add_run(mv.get(key, ""))
                r2.font.name = FONT_BODY
                r2.font.size = Pt(9)
    _section(doc, "2. 关键异动板块解读", _key_movers)

    # 3. 交叉验证
    _section(doc, "3. 交叉验证与资金行为解码",
             lambda: _add_text_block(macro.get("cross_validation", "")))

    # 4. 结论
    def _conclusion():
        concl = macro.get("conclusion", {})
        fields = [
            ("市场阶段", "stage"), ("主攻方向", "attack_direction"),
            ("出逃方向", "retreat_direction"), ("潜在风险", "risks"),
            ("大势预判", "trend"), ("风格定调", "style"), ("核心关注", "watch"),
        ]
        for label, key in fields:
            p = doc.add_paragraph(style="Normal")
            r = p.add_run(f"{label}：")
            r.font.name = FONT_BODY
            r.font.size = Pt(9)
            r.bold = True
            r2 = p.add_run(concl.get(key, ""))
            r2.font.name = FONT_BODY
            r2.font.size = Pt(9)
            if key in ("trend", "style"): r2.bold = True
    _section(doc, "4. 结论与策略前瞻", _conclusion)

    # ── 第三部分（仅周末）──────────────────────────────────────────────────────
    macro_cycle = data.get("macro_cycle")
    if macro_cycle:
        h = doc.add_heading("第三部分 · 宏观周期定位与历史极值映射", level=1)
        for run in h.runs: run.font.name = FONT_BODY
        for sub_title, key in [
            ("1. 当前资产图谱提炼", "current_profile"),
        ]:
            h3 = doc.add_heading(sub_title, level=3)
            for run in h3.runs: run.font.name = FONT_BODY
            _add_text_block(macro_cycle.get(key, ""))
        hist = macro_cycle.get("historical", {})
        h3 = doc.add_heading("2. 历史极值时空定位", level=3)
        for run in h3.runs: run.font.name = FONT_BODY
        for label, key in [("对标年份", "year"), ("标志事件", "event"),
                            ("宏观阶段", "phase"), ("简述", "brief")]:
            p = doc.add_paragraph(style="Normal")
            r = p.add_run(f"{label}：")
            r.font.name = FONT_BODY; r.font.size = Pt(9); r.bold = True
            r2 = p.add_run(hist.get(key, ""))
            r2.font.name = FONT_BODY; r2.font.size = Pt(9)
        for sub_title, key in [("3. 跨时空异同对比", None),
                                ("4. 沙盘推演", None)]:
            h3 = doc.add_heading(sub_title, level=3)
            for run in h3.runs: run.font.name = FONT_BODY
        for label, key in [("形似点", "similarity"), ("神异点", "difference"),
                            ("黑天鹅风险", "risk"), ("隐蔽机会", "opportunity")]:
            p = doc.add_paragraph(style="Normal")
            r = p.add_run(f"{label}：")
            r.font.name = FONT_BODY; r.font.size = Pt(9); r.bold = True
            r2 = p.add_run(macro_cycle.get(key, ""))
            r2.font.name = FONT_BODY; r2.font.size = Pt(9)

    # ── 审计日志 ──────────────────────────────────────────────────────────────
    total   = aligned["total"]
    cat_sum = sum(stats[c]["count"] for c in CAT_ORDER)
    audit   = [
        "【审计日志】",
        f"▸ 合并时段数：{len(snapshots)}",
        f"▸ T0 品种数：{total}",
        f"▸ T0 分类合计校验：{cat_sum} = {total} {'✓' if cat_sum == total else '❌'}",
        f"▸ 跨日追踪表品种数：{len(rating_cands)}",
        "▸ 数据由 Python 确定性计算；分析由 Claude 生成。",
    ]
    p = doc.add_paragraph()
    for line in audit:
        r = p.add_run(line + "\n")
        r.font.name = FONT_MONO
        r.font.size = Pt(8)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"✅ 已写入 {out}")
    return out


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("用法: python render_merge_docx.py <merged.json> <output.docx>")
        sys.exit(1)
    render_merge(sys.argv[1], sys.argv[2])
