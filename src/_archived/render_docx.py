"""
render_docx.py — 把 analyzed JSON 渲染成 Word (.docx) 报告。

所有颜色、字体、列结构固化在本文件，LLM 不参与排版。
表格含「备注」列（默认空白，供用户复盘时手工填写颜色名+批注）。
页面：A4 横向，页边距 0.8cm。
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── 视觉规范（固化，不由 LLM 修改）─────────────────────────────────────────
FONT_BODY = "Microsoft YaHei"
FONT_MONO = "Courier New"

COLOR_UP      = RGBColor(0xFF, 0x00, 0x00)   # 涨：红
COLOR_DOWN    = RGBColor(0x00, 0x00, 0x8B)   # 跌：深蓝
COLOR_ANOMALY = RGBColor(0xFF, 0x00, 0x00)   # **异常** 标记色

HEX_HEADER    = "D9D9D9"   # 表头底色（灰）
HEX_WARN_UP   = "FFE6E6"   # 共振上涨预警背景
HEX_WARN_DOWN = "E6F0FA"   # 共振下跌预警背景

# 列定义 (列名, 宽度/cm)；「备注」列放最末
SINGLE_COLS: list[tuple[str, float]] = [
    ("名称",       2.5),
    ("日期",       2.0),
    ("今涨幅",     1.5),
    ("昨涨幅",     1.5),
    ("涨跌幅差值", 1.8),
    ("今日成交额", 1.5),
    ("成交额环比", 1.8),
    ("归类特征",   1.5),
    ("归类",       1.8),
    ("符合情况",   1.8),
    ("原因分析",   5.5),
    ("备注",       3.0),
]

CAT_ORDER  = ["持续强化", "强反转", "反包修复", "连续杀跌"]
CAT_ROMAN  = {"持续强化": "一", "强反转": "二", "反包修复": "三", "连续杀跌": "四"}


# ── 底层工具 ─────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """设置单元格底色，hex_color 形如 'D9D9D9'（不带 #）。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_row_bg(row, hex_color: str) -> None:
    for cell in row.cells:
        _set_cell_bg(cell, hex_color)


def _cell_write(cell, text: str, *,
                bold: bool = False,
                color: RGBColor | None = None,
                size: int = 9,
                align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空 cell 并写入纯文本（单 run）。"""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text))
    run.font.name = FONT_BODY
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


def _cell_write_analysis(cell, text: str, size: int = 8) -> None:
    """把含 **xxx** / **异常** 标记的分析文本写入 cell，按 run 拆分上色。"""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if not text:
        r = para.add_run("")
        r.font.name = FONT_BODY
        r.font.size = Pt(size)
        return
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        m = re.match(r"\*\*([^*]+)\*\*", part)
        if m:
            inner = m.group(1)
            r = para.add_run(inner)
            r.font.name = FONT_BODY
            r.font.size = Pt(size)
            r.bold = True
            if inner == "异常":
                r.font.color.rgb = COLOR_ANOMALY
        elif part:
            r = para.add_run(part)
            r.font.name = FONT_BODY
            r.font.size = Pt(size)


def _pct_str(v: float) -> str:
    sign = "+" if v > 0 else ""
    return f"{sign}{v * 100:.2f}%"


def _pct_color(v: float) -> RGBColor | None:
    if v > 0: return COLOR_UP
    if v < 0: return COLOR_DOWN
    return None


def _session_date_label(session_label: str, today_date: str) -> str:
    """'0513收盘' + '5月13日' → '2026-05-13(收盘)'。"""
    import datetime as dt, re as _re
    year = dt.datetime.now().year
    m = _re.match(r"(\d+)月(\d+)日", today_date)
    date_iso = f"{year}-{int(m.group(1)):02d}-{int(m.group(2)):02d}" if m else today_date
    suffix = "(中午)" if "中午" in session_label else ("(收盘)" if "收盘" in session_label else "")
    return f"{date_iso}{suffix}"


# ── 文档结构工具 ──────────────────────────────────────────────────────────────

def _setup_landscape(doc: Document) -> None:
    """设置 A4 横向，页边距 0.8cm。"""
    section = doc.sections[0]
    section.orientation  = WD_ORIENT.LANDSCAPE
    section.page_width   = Cm(29.7)
    section.page_height  = Cm(21.0)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Cm(0.8))


def _make_table(doc: Document, col_defs: list[tuple[str, float]]) -> any:
    """创建带灰色表头行的表格，返回 table 对象。"""
    table = doc.add_table(rows=1, cols=len(col_defs))
    table.style = "Table Grid"
    # 列宽
    for i, (_, w) in enumerate(col_defs):
        for cell in table.column_cells(i):
            cell.width = Cm(w)
    # 表头行
    hrow = table.rows[0]
    _set_row_bg(hrow, HEX_HEADER)
    for i, (hdr, _) in enumerate(col_defs):
        _cell_write(hrow.cells[i], hdr, bold=True, size=9)
    return table


def _para(doc: Document, text: str, *,
          bold: bool = False, size: int = 10,
          color: RGBColor | None = None,
          align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """快速添加一段落。"""
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.name  = FONT_BODY
    r.font.size  = Pt(size)
    r.bold       = bold
    if color:
        r.font.color.rgb = color


# ── 主渲染函数 ────────────────────────────────────────────────────────────────

def render_single(analyzed: dict, output_path: str | Path) -> Path:
    """渲染单时段 docx 报告。"""
    doc = Document()
    _setup_landscape(doc)

    session_date = _session_date_label(
        analyzed["session_label"], analyzed["today_date"]
    )
    resonance = analyzed.get("resonance")

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"{analyzed['today_date']} A股全景数据审查报告")
    r.font.name = FONT_BODY
    r.font.size = Pt(18)
    r.bold = True

    # 共振预警
    if resonance == "up":
        _para(doc,
              "🚨 全局极值预警：超70%品种共振起爆！做多情绪极度宣泄，请立刻注意逼空动能！ 🚨",
              bold=True, size=13, color=COLOR_UP, align=WD_ALIGN_PARAGRAPH.CENTER)
    elif resonance == "down":
        _para(doc,
              "⚠️ 全局极值预警：超70%品种共振杀跌！流动性正在急剧收缩，防守退潮信号确立！ ⚠️",
              bold=True, size=13, color=COLOR_DOWN, align=WD_ALIGN_PARAGRAPH.CENTER)

    stats = analyzed["stats"]

    for cat in CAT_ORDER:
        group = analyzed["groups"][cat]
        st    = stats[cat]

        # 分类标题
        heading = doc.add_heading(
            f"{CAT_ROMAN[cat]}、{cat}类（{st['count']}个，占{st['pct']}%）",
            level=2,
        )
        for run in heading.runs:
            run.font.name = FONT_BODY

        table = _make_table(doc, SINGLE_COLS)

        for it in group:
            row   = table.add_row()
            cells = row.cells

            _cell_write(cells[0],  it["name"],           size=9)
            _cell_write(cells[1],  session_date,         size=8)
            _cell_write(cells[2],  _pct_str(it["today_pct"]),
                        color=_pct_color(it["today_pct"]), size=9)
            _cell_write(cells[3],  _pct_str(it["yest_pct"]),
                        color=_pct_color(it["yest_pct"]),  size=9)
            _cell_write(cells[4],  _pct_str(it["pct_diff"]),
                        color=_pct_color(it["pct_diff"]),  size=9)
            _cell_write(cells[5],  it.get("today_amount_str", ""), size=9)
            _cell_write(cells[6],  _pct_str(it["volume_ratio"]),
                        color=_pct_color(it["volume_ratio"]), size=9)
            _cell_write(cells[7],  it.get("feature", ""),     size=9)
            _cell_write(cells[8],  it.get("category", ""),    size=9)
            _cell_write(cells[9],  it.get("compliance", ""),  size=9)
            _cell_write_analysis(cells[10], it.get("analysis", ""), size=8)
            # 备注列：新建报告时默认空白，用户复盘时自行填写
            _cell_write(cells[11], "", size=8, align=WD_ALIGN_PARAGRAPH.LEFT)

        # 统计小注
        _para(doc,
              f"本分类共 {st['count']} 个品种，占总数 {st['pct']}%。",
              size=9)

    # 审计日志
    total   = analyzed["total"]
    cat_sum = sum(stats[c]["count"] for c in CAT_ORDER)
    up_cnt  = analyzed.get("up_count", 0)
    dn_cnt  = analyzed.get("down_count", 0)
    res_txt = ("已触发上涨共振预警" if resonance == "up"
               else "已触发下跌共振预警" if resonance == "down"
               else "未触发（涨跌均不足70%）")
    audit_lines = [
        "【审计日志】",
        f"▸ 总品种数 N = {total}",
        f"▸ 分类合计 = {' + '.join(str(stats[c]['count']) for c in CAT_ORDER)}"
        f" = {cat_sum} {'✓ 闭环' if cat_sum == total else '❌ 不一致'}",
        f"▸ 上涨 {up_cnt} / 下跌 {dn_cnt} / 总数 {total}"
        + (f"（上涨占比 {up_cnt/total*100:.1f}%）" if total else ""),
        f"▸ 极值共振判定：{res_txt}",
        "▸ 本报告所有数字字段由 Python 脚本确定性计算，定性分析由 Claude 生成。",
    ]
    p = doc.add_paragraph()
    for line in audit_lines:
        r = p.add_run(line + "\n")
        r.font.name = FONT_MONO
        r.font.size = Pt(8)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("用法: python render_docx.py <analyzed.json> <output.docx>")
        sys.exit(1)
    data = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
    out  = render_single(data, sys.argv[2])
    print(f"✅ 已写入 {out}")
